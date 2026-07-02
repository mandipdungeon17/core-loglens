"""
Generates the LogLens Invention Disclosure Form (IDF) as a .docx file
following the Honeywell IGS Bangalore Innovation & Ideation template.

Run: python build_idf.py
Output: LogLens-IDF.docx in the project root.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HON_RED = RGBColor(0xE1, 0x00, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_body(doc, text, bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "404040")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(2)
    p_fmt.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E10000")
    pbdr.append(bottom)
    pPr.append(pbdr)


doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ─────────────────────────── TITLE BLOCK ───────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Invention Disclosure Form (IDF)")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = HON_RED

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Innovation & Ideation Flow — Honeywell (IGS Bangalore)")
rs.font.size = Pt(12)
rs.font.color.rgb = GREY

prod = doc.add_paragraph()
prod.alignment = WD_ALIGN_PARAGRAPH.CENTER
rp = prod.add_run("LogLens — Agentless, Index-Free Interactive Log Analytics for Constrained Microservice Environments")
rp.bold = True
rp.font.size = Pt(13)
rp.font.color.rgb = DARK

add_divider(doc)

# Metadata table
meta = add_table(
    doc,
    ["Field", "Details"],
    [
        ["Invention Title", "Request-Adaptive, Index-Free Interactive Querying of Append-Only Time-Ordered Log Files"],
        ["Product / Project", "core-loglens (LogLens) — WES Microservices Log Viewer"],
        ["Business Unit / Site", "Honeywell IGS Bangalore"],
        ["Technology Domain", "Software — Observability / Log Analytics / Distributed Systems Tooling"],
        ["Inventor(s)", "________________________  (to be completed)"],
        ["Contributors", "________________________  (to be completed)"],
        ["Date of First Write-up", "2026-07-02"],
        ["Disclosure Status", "Internal — pre IP review (do not disclose externally)"],
    ],
    col_widths=[1.9, 4.6],
)

doc.add_paragraph()

# ─────────────────────────── STEP 1 ───────────────────────────
add_heading(doc, "Step 1: Problem Statement", level=1)

add_body(doc, "Current situation / existing solution", bold=True)
add_body(
    doc,
    "Honeywell WES (Warehouse Execution System) deployments run a fleet of microservices "
    "(routing, print, momentumconnect, and others — scaling from 3 to 15+ services) directly on "
    "on-premise / edge VMs. Each service continuously appends to large plaintext log files "
    "(commonly 1.5–2.5 GB per service, tens of GB in aggregate). When engineers need to diagnose "
    "an incident, the conventional industry solutions are centralized log platforms — ELK "
    "(Elasticsearch/Logstash/Kibana), Splunk, or a Kafka-to-datastore streaming pipeline — each "
    "of which requires shipping agents, an indexing tier, and a separate database/search cluster.",
)

add_body(doc, "Pain points / limitations of the existing approach", bold=True)
add_bullet(doc, "Heavy infrastructure: ELK/Splunk-class stacks require dedicated indexing nodes, storage, agents, and ongoing operations — infeasible on constrained industrial/edge VMs co-located with the services.", bold_lead="Infrastructure weight — ")
add_bullet(doc, "Log data must be copied/streamed off the host, indexed, and stored again — duplicating tens of GB and adding ingestion latency before logs become searchable.", bold_lead="Data duplication & ingestion lag — ")
add_bullet(doc, "Many deployments are air-gapped or bandwidth-limited; running shipping agents and external clusters is either prohibited or unreliable.", bold_lead="Air-gapped / bandwidth limits — ")
add_bullet(doc, "A naive filesystem grep/tail over a 20+ GB append-only file scans from byte 0 (or issues one syscall per byte via RandomAccessFile.readLine), taking minutes and thrashing memory — unusable interactively.", bold_lead="Naive scanning is too slow — ")
add_bullet(doc, "Fixed thread pools sized as cores×2 either oversubscribe a 4-core VM or exhaust a 512 MB heap when scanning many multi-GB files in parallel.", bold_lead="Resource fragility — ")
add_bullet(doc, "Re-running the same investigation with a slightly changed filter forces a full re-scan; no cheap way to refine a query interactively.", bold_lead="No cheap query refinement — ")

add_body(doc, "Why this is worth solving (impact)", bold=True)
add_bullet(doc, "Faster mean-time-to-diagnosis for WES production incidents directly reduces warehouse downtime and SLA penalties.")
add_bullet(doc, "Eliminating ELK/Splunk infrastructure removes licensing, hardware, and operational cost for every deployment site.")
add_bullet(doc, "A single self-contained tool that runs anywhere the services run is deployable to air-gapped and edge sites where centralized observability cannot go.")

add_divider(doc)

# ─────────────────────────── STEP 2 ───────────────────────────
add_heading(doc, "Step 2: Desired Solution", level=1)

add_body(doc, "What we want to achieve (success criteria)", bold=True)
add_bullet(doc, "Interactive (sub-second to few-second) querying over tens of GB of unindexed, append-only logs from a single executable, with no external index, database, agent, or message bus.")
add_bullet(doc, "Sub-second response for query refinement and pagination on an already-scanned result set (no repeated disk I/O).")
add_bullet(doc, "Correct results for both recent (‘last 5 minutes’) and historical (‘14:00–15:00 yesterday’) time windows without scanning the whole file.")
add_bullet(doc, "Elastic footprint: run safely on a 4-core / 512 MB VM for 3 services and scale to a 16-core / 4 GB host for 15+ services with the same binary.")

add_body(doc, "Key performance / functional requirements", bold=True)
add_bullet(doc, "Structural filters (level, traceId, spanId, userId, siteId, tenantId, logger, message, time range) plus a free-form boolean query language (AND/OR/NOT, field:value, quoted phrases, grouping).")
add_bullet(doc, "Multi-user concurrency with per-user isolation and bounded memory.")
add_bullet(doc, "Support for rotated and gzip-compressed archive files.")
add_bullet(doc, "Deterministic memory ceiling regardless of input volume (no OutOfMemory under concurrent load).")

add_body(doc, "Constraints (cost, space, compatibility, safety)", bold=True)
add_bullet(doc, "Cost: zero additional infrastructure/licensing; single JAR.")
add_bullet(doc, "Space/compute: must operate within a small heap (as low as 512 MB) alongside the monitored services.")
add_bullet(doc, "Compatibility: read existing WES log formats unchanged (5 on-disk formats); no change to how services log.")
add_bullet(doc, "Safety/Security: read-only on log files; path-traversal-safe historical browsing; internal diagnostic tool with network-level access control.")

add_divider(doc)

# ─────────────────────────── STEP 3 ───────────────────────────
add_heading(doc, "Step 3: Idea / Innovation Description", level=1)

add_body(
    doc,
    "LogLens is a single Spring Boot (Java 21) executable that reads raw log files directly off "
    "the local filesystem and makes them interactively queryable — with no index, no database, no "
    "agent, and no streaming pipeline. It exploits the fact that log files are append-only and "
    "time-ordered to avoid ever building an index: instead of pre-indexing, it decides at query "
    "time HOW to read each file (direction + I/O method), uses timestamp binary search to jump "
    "straight to the relevant byte range, and caches structural results so that query refinements "
    "and pagination cost zero disk I/O.",
)

add_body(doc, "What changes versus the current design", bold=True)
add_table(
    doc,
    ["Aspect", "Baseline (ELK/Splunk/grep)", "LogLens"],
    [
        ["Indexing", "Pre-built inverted index / datastore", "None — query-time adaptive scan of raw files"],
        ["Infrastructure", "Agents + index tier + DB cluster", "Single self-contained JAR on the host"],
        ["Time-window lookup", "Index seek, or full file scan (grep)", "Timestamp binary search over raw text"],
        ["Read method", "Fixed / one-syscall-per-byte", "Per-request choice of direction + bulk I/O"],
        ["Query refinement", "Re-query the index / re-scan", "Zero-I/O refilter from in-memory cache"],
        ["Parallelism", "Cluster / fixed pool", "CPU- and heap-bounded elastic pool"],
        ["Memory", "Offloaded to cluster", "Deterministic bounded-merge ceiling (100K)"],
    ],
    col_widths=[1.5, 2.6, 2.4],
)

add_body(doc, "System overview (block diagram)", bold=True, space_after=2)
overview = (
    "  ┌──────────────────────────────────────────────────────────────────────┐\n"
    "  │                        LogLens (single JAR)                            │\n"
    "  │                                                                        │\n"
    "  │   Browser SPA  ──HTTP──►  REST Controller                              │\n"
    "  │   (virtual scroll)                │                                    │\n"
    "  │                                   ▼                                    │\n"
    "  │                        ┌─────────────────────┐                         │\n"
    "  │                        │  Strategy Selector  │  {sort, time?, filters, │\n"
    "  │                        │ (per file/request)  │   .gz?} → strategy      │\n"
    "  │                        └─────────┬───────────┘                         │\n"
    "  │        BACKWARD / FORWARD / BINARY_THEN_BACKWARD / BINARY_THEN_FORWARD  │\n"
    "  │                                  ▼                                     │\n"
    "  │   Timestamp Binary Search ─► Bulk-I/O Scan (parallel, bounded pool)     │\n"
    "  │                                  │  Pass 1: structural filter + early-exit│\n"
    "  │                                  ▼                                     │\n"
    "  │              Bounded PriorityQueue Merge (100K cap)                     │\n"
    "  │                                  │  Pass 2: boolean query (post/pre-cap) │\n"
    "  │                                  ▼                                     │\n"
    "  │        Two-Layer Cache:  L1 structural (by fingerprint)                │\n"
    "  │                          L2 per-session query view + cursor            │\n"
    "  │                                  ▼                                     │\n"
    "  │                     JSON results + searchId  ──► Browser               │\n"
    "  └──────────────────────────────────────────────────────────────────────┘\n"
    "                                    ▲\n"
    "        Reads directly ────────────┘\n"
    "        Local filesystem: service-1.log, service-2.log … *.log.gz (append-only)"
)
mono = doc.add_paragraph()
mrun = mono.add_run(overview)
mrun.font.name = "Consolas"
mrun.font.size = Pt(8)

add_divider(doc)

# ─────────────────────────── STEP 4 ───────────────────────────
add_heading(doc, "Step 4: Working Principle", level=1)

add_body(doc, "Step-by-step operation", bold=True)
add_numbered(doc, "A search request arrives with sort order (ASC/DESC), optional time window, structural filters, an optional free-form query, and scan depth.", bold_lead="Request intake — ")
add_numbered(doc, "If the request carries a valid searchId, the next page is served from the in-memory cache cursor with zero disk I/O and the flow ends.", bold_lead="Cache short-circuit — ")
add_numbered(doc, "A filter fingerprint is computed; if a fresh structural cache exists (validated against live file sizes, since append-only files only grow), it is reused and only the query view is (re)derived.", bold_lead="Cache reuse check — ")
add_numbered(doc, "For each target file, a strategy is chosen from {sortOrder, hasTimeFilter, hasStructuralFilters, isGzip}: BACKWARD (bulk 64 KB chunks from EOF), FORWARD (buffered read from start), BINARY_THEN_BACKWARD, or BINARY_THEN_FORWARD.", bold_lead="Per-file strategy selection — ")
add_numbered(doc, "For time-windowed queries, a binary search probes timestamps (seek→read complete line→compare) to converge (~20 iterations, capped at 64) on the byte offset of the window boundary, adding a 1 MB safety margin so no boundary entry is cut.", bold_lead="Timestamp binary search — ")
add_numbered(doc, "Files are scanned in parallel on a resource-aware pool. During the scan (Pass 1) each parsed entry is tested against structural filters; time filters trigger early-exit that stops reading once the scan leaves the window.", bold_lead="Parallel bounded scan + Pass 1 — ")
add_numbered(doc, "Per-service results feed a bounded PriorityQueue (100K cap) that keeps the newest (DESC) or oldest (ASC) entries and progressively releases memory.", bold_lead="Bounded merge — ")
add_numbered(doc, "The free-form boolean query (Pass 2) is applied AFTER the merge at default depth (so it can be re-run from cache with no I/O), or BEFORE the cap for deep scans (so rare deep matches are not evicted). The choice is recorded in the cache fingerprint and a ‘bakedQuery’ marker.", bold_lead="Depth-adaptive Pass 2 — ")
add_numbered(doc, "Results plus a searchId are returned; the browser renders only ~80 DOM rows via virtual scroll and requests further pages by searchId.", bold_lead="Response & pagination — ")

add_body(doc, "Key components involved", bold=True)
add_bullet(doc, "Strategy Selector — maps query semantics to read direction + I/O method.")
add_bullet(doc, "Timestamp Binary Search — index-free byte-offset locator over raw time-ordered text.")
add_bullet(doc, "Two-Pass Filter (structural in-scan + boolean query) with depth-adaptive ordering.")
add_bullet(doc, "Two-Layer Cache (L1 structural by fingerprint, L2 per-session query view + cursor).")
add_bullet(doc, "Resource-Aware Thread Pool and Bounded PriorityQueue Merge.")
add_bullet(doc, "Log Parser (5 on-disk formats) and recursive-descent boolean Query Engine.")

add_body(doc, "Control / decision logic", bold=True)
add_bullet(doc, "Strategy: .gz→FORWARD; ASC+time→BINARY_THEN_FORWARD; ASC+structural(no time)→BACKWARD; ASC+none→FORWARD; DESC+time→BINARY_THEN_BACKWARD; else→BACKWARD.")
add_bullet(doc, "Pool size = min(serviceCount, min(cores×0.6, heapMB/150)); overridable by config.")
add_bullet(doc, "Query ordering: if resolveMaxScan(req) > defaultMaxScan → apply query pre-cap (hybrid, ‘baked’); else post-cap (cacheable).")
add_bullet(doc, "Cache freshness: reuse L1 iff stored per-file sizes == current sizes; sliding 15-min TTL; LRU eviction at 5 caches / 10 sessions each.")

add_divider(doc)

# ─────────────────────────── STEP 5 ───────────────────────────
add_heading(doc, "Step 5: Novelty / Novelties", level=1)

add_body(doc, "What is new or unique", bold=True)
add_bullet(doc, "A request-adaptive engine that selects both READ DIRECTION and I/O METHOD per file per query from query semantics, then binary-searches raw (unindexed) time-ordered logs to jump directly into the target time window and scan INTO it (backward from toTime for DESC, forward from fromTime for ASC).", bold_lead="N1 — ")
add_bullet(doc, "A depth-adaptive ‘hybrid query’ that REORDERS free-form query filtering relative to a bounded merge — post-cap at default depth (cacheable, zero-I/O refilter) and pre-cap for deep scans (maximizing recall of rare deep matches) — with the ordering decision ENCODED into the cache fingerprint and a ‘bakedQuery’ marker so the cache layer stays correct.", bold_lead="N2 — ")
add_bullet(doc, "A two-layer cache that separates STRUCTURAL results (Layer 1, shared by fingerprint, freshness-validated against live append-only file sizes) from PER-SESSION query views with independent cursors (Layer 2), so a changed query or ‘Load More’ costs zero disk I/O.", bold_lead="N3 — ")
add_bullet(doc, "Dual-bounded elastic parallelism keyed to LIVE workload — parallel scan width simultaneously bounded by CPU, available heap, and the actual number of services with existing files.", bold_lead="N4 — ")
add_bullet(doc, "The integrated agentless, index-free architecture that combines N1–N4 to make tens of GB of raw logs interactively queryable from one co-located JAR.", bold_lead="N5 — ")

add_body(doc, "Comparison with existing solutions", bold=True)
add_table(
    doc,
    ["Capability", "ELK / Splunk", "grep / tail / RAF", "LogLens"],
    [
        ["Index required", "Yes", "No", "No"],
        ["External infra", "Yes (cluster)", "No", "No"],
        ["Time-window jump", "Index seek", "Full scan", "Timestamp binary search on raw text"],
        ["Adaptive read direction/method", "N/A", "No", "Yes (per file/request)"],
        ["Zero-I/O query refine", "No (re-query)", "No", "Yes (L1→L2 refilter)"],
        ["Depth-adaptive recall/latency", "N/A", "No", "Yes (baked vs cacheable)"],
        ["Bounded memory guarantee", "Cluster-side", "No", "Yes (100K bounded merge)"],
    ],
    col_widths=[1.9, 1.4, 1.3, 2.0],
)

add_body(doc, "Why competitors do not have this today", bold=True)
add_bullet(doc, "The mainstream mental model is ‘index first, then query’; the counter-intuitive step here is deciding read strategy at query time and searching raw files with no index at all.")
add_bullet(doc, "Adaptive reordering of the query relative to the memory cap — trading recall against cacheable latency and reflecting that choice in cache identity — is a non-obvious control decision most log tools never face because they delegate to a search cluster.")
add_bullet(doc, "The design is driven by an unusual constraint (interactive querying on constrained, co-located, possibly air-gapped edge VMs) that centralized-observability vendors do not target.")

add_divider(doc)

# ─────────────────────────── STEP 6 ───────────────────────────
add_heading(doc, "Step 6: Novelty Drafting Guidance (Applied)", level=1)

add_body(doc, "Novelty type classification", bold=True)
add_bullet(doc, "New mechanism: query-time strategy selector + timestamp binary search over unindexed logs.")
add_bullet(doc, "New combination: strategy selection + bounded merge + two-layer cache + depth-adaptive query ordering working as one pipeline.")
add_bullet(doc, "Simplification: removes the entire index/agent/datastore tier (order-of-magnitude infrastructure reduction).")
add_bullet(doc, "Performance jump: 43–111× faster query refinement via cache; ~0.014 s paginated ‘Load More’; 6.6× faster time-range lookup via binary search; 52% fewer humongous GC allocations and 15× faster sort after bounded merge.")

add_body(doc, "Novelty statements (explicit, technical, tied to advantage)", bold=True)
add_body(
    doc,
    "\u201cA query-time strategy selector chooses both read direction and I/O method per log file and "
    "uses timestamp binary search over unindexed, append-only files to seek directly into the "
    "requested time window — enabling interactive time-range queries over tens of GB without any "
    "index.\u201d",
    italic=True,
)
add_body(
    doc,
    "\u201cA depth-adaptive query stage reorders free-form filtering relative to a bounded-memory "
    "merge — post-cap for cacheable zero-I/O refinement at default depth, pre-cap for maximal "
    "recall at deep scan depth — and encodes this ordering into the cache key so refinement remains "
    "correct and instantaneous.\u201d",
    italic=True,
)

add_divider(doc)

# ─────────────────────────── STEP 7 ───────────────────────────
add_heading(doc, "Step 7: Claims & Summary Claim", level=1)

add_body(doc, "Claim 1 — Core functional claim", bold=True)
add_body(
    doc,
    "A method for interactively querying unindexed, append-only, time-ordered log files, "
    "comprising: receiving a query specifying a sort order, an optional time window, structural "
    "filters, and an optional free-form boolean expression; selecting, for each file and per "
    "request, a scan strategy that determines both a read direction and an I/O method from the "
    "query semantics and file type; when a time window is specified, performing a timestamp binary "
    "search that seeks to and reads complete lines at successive byte offsets to converge on a "
    "boundary offset and then scans into the window; applying structural filters during the scan "
    "with time-based early termination; merging per-file results through a fixed-capacity priority "
    "queue that retains entries by sort order; and returning results without use of any pre-built "
    "index, external database, log-shipping agent, or message bus.",
)

add_body(doc, "Claim 2 — Performance / cost advantage claim", bold=True)
add_body(
    doc,
    "The method of Claim 1, further comprising caching the structural scan result in a first layer "
    "keyed by a filter fingerprint and validated for freshness against current file sizes, and "
    "deriving per-session query views with independent pagination cursors in a second layer, such "
    "that a change to the free-form expression or a request for a subsequent page is served from "
    "memory with zero disk input/output; and executing the scan on a thread pool whose width is "
    "bounded simultaneously by processor count, available heap, and the number of files present — "
    "thereby delivering interactive latency (order 0.01–0.4 s for refinement/pagination) on tens "
    "of gigabytes of logs on a single constrained host without additional infrastructure.",
)

add_body(doc, "Summary claim (one-line value statement)", bold=True)
add_body(
    doc,
    "\u201cLogLens makes tens of gigabytes of raw, unindexed service logs interactively searchable "
    "from a single co-located executable — no index, no agent, no database — by choosing how to "
    "read each file at query time and caching results for instant, zero-I/O refinement.\u201d",
    italic=True,
    bold=True,
)

add_divider(doc)

# ─────────────────────────── STEP 8 ───────────────────────────
add_heading(doc, "Step 8: Claim & Sub-Claim Drafting Guidance (Applied)", level=1)

add_body(doc, "Main claim (functional, single sentence)", bold=True)
add_body(
    doc,
    "A system that queries unindexed, append-only time-ordered log files by selecting a per-file "
    "read direction and I/O method from the query, locating time windows via timestamp binary "
    "search, filtering during the scan, and merging through a fixed-capacity priority queue — "
    "delivering interactive results with no index, agent, or database.",
    italic=True,
)

add_body(doc, "Sub-claims (dependent, broadening coverage)", bold=True)
add_numbered(doc, "wherein the scan strategy is selected from {backward bulk-chunk, forward buffered, binary-then-backward, binary-then-forward} based on sort order, presence of a time filter, presence of structural filters, and whether the file is gzip-compressed.")
add_numbered(doc, "wherein the timestamp binary search skips a partial line after each seek, reads up to a bounded number of lines to find a parseable timestamp, caps iterations at log2(fileSize), and applies a fixed byte safety margin so boundary entries are not truncated.")
add_numbered(doc, "wherein a free-form boolean filtering stage is applied after the priority-queue cap at a default scan depth and before the cap when the requested scan depth exceeds the default, the ordering being encoded in the cache key.")
add_numbered(doc, "wherein a first cache layer stores structural results keyed by a filter fingerprint and validated against current file sizes, and a second cache layer derives per-user query views with independent pagination cursors, serving refinements and additional pages with zero disk I/O.")
add_numbered(doc, "wherein the parallel scan width is computed as the minimum of the file count, a fraction of processor cores, and available heap divided by a per-thread memory budget.")
add_numbered(doc, "wherein rotated and gzip-compressed archive files are supported, gzip files being decompressed once to a reusable temporary file for random-access scanning.")

add_divider(doc)

# ─────────────────────────── STEP 9 ───────────────────────────
add_heading(doc, "Step 9: Value for Business", level=1)

add_body(doc, "Customer value", bold=True)
add_bullet(doc, "Faster incident diagnosis → less WES/warehouse downtime and fewer SLA breaches.")
add_bullet(doc, "Works on constrained, edge, and air-gapped sites where centralized observability cannot be deployed.")
add_bullet(doc, "Zero change to how services log and zero learning curve for a heavyweight platform.")

add_body(doc, "Business value", bold=True)
add_bullet(doc, "Eliminates ELK/Splunk licensing, hardware, and operations cost per site (COGS/TCO reduction).")
add_bullet(doc, "Differentiator for Honeywell WES offerings: built-in, no-infrastructure diagnostics.")
add_bullet(doc, "Single ~single-JAR deliverable → trivial packaging, patching, and support.")

add_body(doc, "Scalability across products / platforms", bold=True)
add_bullet(doc, "Same binary scales 3→15+ services and 512 MB→4 GB heaps via resource-aware pooling.")
add_bullet(doc, "Format-agnostic parser (5 formats) and config-only service onboarding — reusable for any Honeywell Java/microservice product line that writes time-ordered logs.")

add_body(doc, "Quantified evidence (from load testing)", bold=True)
add_table(
    doc,
    ["Metric", "Result"],
    [
        ["Test scale", "13 services × 2.5 GB (≈31 GB), 22-core / -Xmx4g"],
        ["Fresh scan (typical)", "1–8 s depending on services/filter"],
        ["Cache-hit query refinement", "0.02–0.4 s (43–111× faster than rescan)"],
        ["Paginated ‘Load More’", "≈0.014 s (from cache cursor)"],
        ["Time-range binary search", "0.5 s vs 3.3 s naive (6.6× faster)"],
        ["Bounded merge GC impact", "Humongous allocations 685→330 (−52%); sort ~15× faster"],
        ["Stability under 5 concurrent users", "Peak heap ≈3.3 GB, 0 Full GCs"],
    ],
    col_widths=[2.6, 3.9],
)

add_divider(doc)

# ─────────────────────────── STEP 10 ───────────────────────────
add_heading(doc, "Step 10: Risks & Mitigation", level=1)

add_table(
    doc,
    ["Risk", "Type", "Mitigation", "Validation"],
    [
        ["Prior art on binary search over log files / log caching individually", "IP", "Claim the specific ADAPTIVE ORCHESTRATION (N1+N2+N3), not individual parts; emphasize depth-adaptive query ordering encoded in cache key", "Prior-art / FTO search before filing"],
        ["Non-monotonic timestamps or malformed lines break binary search", "Technical", "Iteration cap (log2), 1 MB safety margins, multi-line timestamp probing, fallback to full scan", "Unit tests on edge cases; fuzz malformed logs"],
        ["Very rare matches deep in huge files missed by 100K cap", "Technical", "Depth-adaptive hybrid query applies filter pre-cap on deep scans", "Deep-scan recall tests (already covered)"],
        ["Memory pressure under many concurrent users", "Technical", "Bounded merge (100K), 5-cache/10-session LRU, sliding TTL, dual-bounded pool", "Concurrency load tests; GC logging"],
        ["Log format drift across products", "Commercial", "Pluggable multi-format parser; config-only onboarding", "Parser regression suite (5 formats)"],
        ["Security (path traversal, log injection, XSS)", "Technical/Security", "Canonical-path validation, CRLF sanitization, output encoding", "Coverity/BlackDuck remediation completed"],
    ],
    col_widths=[1.9, 0.9, 2.1, 1.6],
)

add_body(doc, "Validation plan", bold=True, space_after=2)
add_bullet(doc, "Simulation/benchmark: repeatable load_test.py + stress_metrics.py suites with GC logging on representative 13–15 service datasets.")
add_bullet(doc, "Prototype: current working implementation (114 passing tests) serves as the proof of concept.")
add_bullet(doc, "Field pilot: deploy alongside a live WES site and measure MTTR improvement vs baseline tooling.")

add_divider(doc)

# ─────────────────────────── STEP 11 ───────────────────────────
add_heading(doc, "Step 11: Next Steps", level=1)
add_numbered(doc, "Proof of concept: mature working implementation already exists; package a demo dataset + scripted scenarios for reviewers.", bold_lead="PoC — ")
add_numbered(doc, "Feasibility & validation: run field pilot at 1–2 WES sites; capture MTTR, resource, and reliability metrics vs current tooling.", bold_lead="Feasibility — ")
add_numbered(doc, "IP review: conduct prior-art/FTO search focused on adaptive query-time strategy selection + depth-adaptive cache-encoded query ordering; keep disclosure internal until cleared.", bold_lead="IP — ")
add_numbered(doc, "Submission: file via the Honeywell Innovation / VE portal with this IDF, claims, and quantified evidence.", bold_lead="Submit — ")

add_divider(doc)

# ─────────────────────────── CHECKLIST ───────────────────────────
add_heading(doc, "Checklist (Ready for Reviews)", level=1)
add_table(
    doc,
    ["Done", "Section", "Quality check"],
    [
        ["\u2611", "Problem Statement", "Current situation, pain points, and impact are clear."],
        ["\u2611", "Desired Solution", "Success criteria, requirements, and constraints stated."],
        ["\u2611", "Idea Description", "Concept and deltas vs baseline clear; diagram included."],
        ["\u2611", "Working Principle", "Step-by-step operation and key components explained."],
        ["\u2611", "Novelty Drafting", "Mechanism / combination / simplification / performance jump captured."],
        ["\u2611", "Novelty Wording", "Differentiation explicit, technical, tied to advantage."],
        ["\u2611", "Main Claim", "One-sentence claim: elements + relationship + outcome."],
        ["\u2611", "Sub-Claims", "6 sub-claims cover variations, logically dependent."],
        ["\u2611", "Value for Business", "Customer + business value quantified; scalability stated."],
        ["\u2611", "Risks & Mitigation", "Key risks, mitigations, and validation plan included."],
        ["\u2611", "Next Steps", "PoC / feasibility / validation / submission path defined."],
    ],
    col_widths=[0.6, 1.7, 4.2],
)

add_divider(doc)

# ─────────────────────────── PATENTABILITY ───────────────────────────
add_heading(doc, "Patentability Considerations (Pre-check)", level=1)

add_body(doc, "1. Novelty vs prior art", bold=True)
add_bullet(doc, "Exact technical difference: query-time selection of read direction + I/O method AND depth-adaptive query ordering encoded in the cache key — over UNINDEXED append-only files. Core novelty + variants captured in sub-claims 1–6.")
add_body(doc, "2. Inventive step / non-obviousness", bold=True)
add_bullet(doc, "A skilled engineer’s default is to build an index; deliberately not indexing and instead adapting the read at query time — plus reordering the query around a memory cap and reflecting that in cache identity — overcomes a recall-vs-latency trade-off with a measurable performance jump (43–111×).")
add_body(doc, "3. Enablement", bold=True)
add_bullet(doc, "Working principle (Step 4) plus the existing implementation and 114 tests demonstrate it can be built. Critical parameters: 64 KB backward chunk, 1 MB binary-search margin, 200-line timestamp probe, 100K merge cap, pool = min(services, cores×0.6, heapMB/150), 15-min TTL.")
add_body(doc, "4. Claim strategy", bold=True)
add_bullet(doc, "Main claim = essential elements + relationship + outcome (Step 7/8). Sub-claims cover strategy set, binary-search robustness, depth-adaptive ordering, two-layer cache, dual-bounded pool, gzip/rotation handling.")
add_body(doc, "5. Disclosure hygiene", bold=True)
add_bullet(doc, "Keep internal until IP review; dated first write-up 2026-07-02; record contributors. Avoid external/public disclosure if patenting is intended.")

# Footer note
doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run(
    "Prepared for internal Honeywell IGS Bangalore innovation review. Evidence sourced from the "
    "core-loglens codebase (LogSearchService, LogParserService, QueryEngine) and project "
    "documentation (ARCHITECTURE.md, STRATEGY.md, Technical.md, README.md, TODO-scalability.md)."
)
nr.italic = True
nr.font.size = Pt(9)
nr.font.color.rgb = GREY

doc.save("LogLens-IDF.docx")
print("Saved LogLens-IDF.docx")
